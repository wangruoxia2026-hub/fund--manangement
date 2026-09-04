# -*- coding: utf-8 -*-
"""
基金合同管理系统 - Flask 后端 API
提供基金、节点、提取信息、文件的增删改查接口
支持 DeepSeek AI 智能提取功能
"""

import os
import re
import json
import sqlite3
import base64
from datetime import datetime
from flask import Flask, request, jsonify, g, send_from_directory, render_template
from werkzeug.utils import secure_filename

# DeepSeek API 配置
DEEPSEEK_API_KEY = "sk-526e6d39be824a1982f935024f78e2b7"
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"

# Flask 应用配置
app = Flask(__name__, template_folder='templates')
app.config['JSON_AS_ASCII'] = False  # 支持中文 JSON
app.config['UPLOAD_FOLDER'] = 'uploads'  # 文件上传目录
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 最大上传 16MB
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'doc', 'docx', 'jpg', 'png', 'jpeg'}

# 数据库配置
DATABASE = 'fund_management.db'


# ==================== 数据库连接工具 ====================

def get_connection():
    """获取数据库连接"""
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn


def dict_from_row(row):
    """将 sqlite3.Row 转换为字典"""
    if row is None:
        return None
    return dict(zip(row.keys(), row))


# ==================== DeepSeek AI 提取功能 ====================

def extract_with_deepseek(content, node_type):
    """
    使用 DeepSeek API 提取合同信息
    
    Args:
        content: 合同文本内容
        node_type: 节点类型 (首次/变更/分红/临开/自定义)
    
    Returns:
        dict: 提取的信息
    """
    import urllib.request
    import urllib.error
    
    # 检查内容是否有效
    if not content or len(content.strip()) < 50:
        raise ValueError("合同文本内容过短或为空，无法提取信息")
    
    # 检查是否全是乱码/非中文内容
    chinese_count = sum(1 for c in content if '\u4e00' <= c <= '\u9fff')
    if chinese_count < 50:
        raise ValueError("合同文本中中文字符过少，可能存在编码问题")
    
    # 根据节点类型构建不同的提取提示
    # 关键字段定位提示：告诉模型字段通常出现在合同的哪些章节
    field_hints = """
【字段定位提示（请优先在这些章节寻找）】
- product_name 通常在合同封面/前言部分（如"基金合同"、"私募基金合同"），有时在合同主体第一段或"基金基本情况"章节
- custodian（基金托管人）通常在合同首页"基金管理人：XXX、基金托管人：XXX"处，或"基金的基本情况"章节
- establish_date（成立日期）通常在"基金的基本情况"、"基金的成立与备案"章节
- investment_manager（投资经理/投资主办人）通常在"基金管理人"、"基金的投资管理"章节的描述中（人名）
- strategy（投资策略）通常在"基金的投资范围和投资策略"、"基金的投资"章节
- subscription_fee（认申购费）通常在"基金的申购与赎回"、"基金的费用与税收"章节
- custody_fee（托管费+运营外包费）通常在"基金的费用与税收"、"基金费用"章节
- risk_return（风险收益特征）通常在"基金的风险揭示"、"基金的风险评级"章节
- lock_period（锁定期/封闭期）通常在"基金的申购与赎回"、"基金份额的转让"章节
- annual_fee（年管理费）通常在"基金的费用与税收"、"基金费用"章节
- performance_fee（业绩报酬）通常在"基金的费用与税收"、"业绩报酬"章节
- investment_scope（投资范围）通常在"基金的投资范围"、"基金的投资"章节
- investment_ratio（投资比例/资产配置比例）通常在"基金的投资组合"、"投资限制"章节
"""

    if node_type in ['首次', '变更']:
        system_prompt = f"""你是一个专业的私募基金合同信息提取助手。你的任务是从给定的合同文本中，准确、完整地提取关键信息。{field_hints}

【严格要求】
1. 必须返回严格的 JSON 对象，不要包含任何 Markdown 标记、解释、说明文字。
2. 找不到的字段填空字符串 ""，不要填 null，不要编造。
3. 中文日期统一转换为 YYYY-MM-DD 格式（例如 "2024年1月15日" → "2024-01-15"，"2024年1月" → "2024-01"，"不定期" → ""）。
4. 费率类字段保留原文表述，例如 "1.5%"、"0.25%/年"。
5. 如果字段在合同中未明确说明（例如"不收取认购费"），也要如实填入"无"或"不收取"。
6. 如果同一个字段在合同中多次出现，提取最权威的表述（合同正文 > 附则 > 招募说明书摘要）。
7. 必须通读全文后再提取，不要只看开头部分。

【需要提取的字段及说明】
- product_name: 基金/产品的完整中文全称（注意：基金名称中的括号内容、空格要保留）
- custodian: 基金托管人全称（通常是银行或证券公司，例如"招商证券股份有限公司"）
- establish_date: 基金成立日期，YYYY-MM-DD 格式
- investment_manager: 投资经理姓名（或投资主办人、基金经理，通常是人名）
- strategy: 投资策略的实质性描述（例如"股票多头"、"量化对冲"、"CTA策略"、"固收+相对价值"等，不要填"控制风险"这种笼统表述）
- subscription_fee: 认购费/申购费费率（保留原文，例如"1.5%"、"无"、"不收取"）
- custody_fee: 托管费 + 运营外包服务费（合并表述，例如"0.25%/年"）
- risk_return: 基金风险等级与收益特征（例如"R3 中等风险"、"中等风险"）
- lock_period: 锁定期/封闭期（保留原文表述，例如"6个月"、"12个月"、"不定期"）
- annual_fee: 管理费（年管理费），例如"1.5%/年"
- performance_fee: 业绩报酬提取比例及提取方式，例如"提取20%"，无则填"无"
- investment_scope: 投资范围（详细说明，包括可投资品种）
- investment_ratio: 投资比例/资产配置比例（例如"股票 80%-95%, 现金 5%-20%"）

【输出格式】
{{"product_name":"","custodian":"","establish_date":"","investment_manager":"","strategy":"","subscription_fee":"","custody_fee":"","risk_return":"","lock_period":"","annual_fee":"","performance_fee":"","investment_scope":"","investment_ratio":""}}"""

        user_prompt = f"请从以下基金合同全文中提取上述所有字段。合同较长，请通读全文后再严格按照 JSON 格式返回：\n\n{content}"

    elif node_type == '分红':
        system_prompt = """你是一个专业的基金分红信息提取助手。请从给定的文本中准确提取分红相关信息。

【严格要求】
1. 必须返回严格的 JSON 对象，不要包含任何 Markdown 标记或解释。
2. 找不到的字段填空字符串 ""，不要填 null，不要编造。
3. 日期统一转换为 YYYY-MM-DD 格式。

【需要提取的字段】
- dividend_count: 分红次数（数字字符串，例如 "1"、"3"）
- dividend_date: 分红日期 / 权益登记日，YYYY-MM-DD 格式
- dividend_amount: 分红金额（每份或总额，保留数字和单位，例如 "0.05元/份" 或 "500万元"）

【输出格式】
{"dividend_count":"","dividend_date":"","dividend_amount":""}"""

        user_prompt = f"请从以下文本中提取分红信息，严格按 JSON 格式返回。请通读全文后再提取：\n\n{content}"

    else:
        # 临开、自定义节点只提取备注
        system_prompt = """你是一个专业的基金信息提取助手。请从给定的文本中提取关键备注信息。

【严格要求】
1. 必须返回严格的 JSON 对象。
2. 备注字段要尽量完整地保留重要信息，包括但不限于：会议时间、地点、参与方、决议事项、关键数据、变更说明等。

【输出格式】
{"remark":""}"""

        user_prompt = f"请从以下文本中提取备注信息，严格按 JSON 格式返回。请通读全文后再提取：\n\n{content}"

    # 调用 DeepSeek API（开启 JSON 模式 + 增加 max_tokens）
    payload = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": 0.1,
        "max_tokens": 8192,
        "response_format": {"type": "json_object"}
    }
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
    }
    
    try:
        req = urllib.request.Request(
            DEEPSEEK_API_URL,
            data=json.dumps(payload).encode('utf-8'),
            headers=headers,
            method='POST'
        )

        with urllib.request.urlopen(req, timeout=120) as response:
            result = json.loads(response.read().decode('utf-8'))
            content_text = result['choices'][0]['message']['content']

            # 打印原始响应，方便排查问题
            print(f"[DEBUG] DeepSeek raw response (first 800 chars): {content_text[:800]}")

            # 优先尝试直接解析（response_format=json_object 时直接是合法 JSON）
            try:
                return json.loads(content_text)
            except json.JSONDecodeError:
                # 兜底：从返回中截取 JSON 部分
                json_match = re.search(r'\{[\s\S]*\}', content_text)
                if json_match:
                    return json.loads(json_match.group())
                else:
                    print(f"[WARNING] DeepSeek response doesn't contain JSON")
                    return {}
    except urllib.error.HTTPError as e:
        error_body = e.read().decode('utf-8') if e.fp else str(e)
        print(f"[ERROR] DeepSeek API HTTP error {e.code}: {error_body}")
        raise ValueError(f"DeepSeek API 请求失败 (HTTP {e.code}): {error_body}")
    except urllib.error.URLError as e:
        print(f"[ERROR] DeepSeek API connection error: {e.reason}")
        raise ValueError(f"无法连接 DeepSeek API: {e.reason}")
    except json.JSONDecodeError as e:
        print(f"[ERROR] Failed to parse DeepSeek response: {e}")
        raise ValueError("DeepSeek 返回格式错误")


def read_file_content(file_path):
    """读取文件内容"""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        # 优先使用 pdfplumber（中文合同解析更稳定），兼容 PyPDF2
        text = ''
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                # 最多取前 10 页，避免内容过长
                pages = pdf.pages[:10]
                parts = []
                for page in pages:
                    page_text = page.extract_text() or ''
                    if page_text:
                        parts.append(page_text)
                text = '\n'.join(parts)
        except ImportError:
            # 兼容旧方案
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    parts = []
                    for page in reader.pages[:10]:
                        parts.append(page.extract_text() or '')
                    text = '\n'.join(parts)
            except ImportError:
                return "【PDF文件需要安装 pdfplumber 或 PyPDF2 库才能解析，请运行: pip install pdfplumber】"
        except Exception as e:
            print(f"[ERROR] Failed to read PDF {file_path}: {e}")
            return f"【无法解析PDF文件: {str(e)}，请手动输入合同信息】"

        if not text or not text.strip():
            return "【PDF文件解析为空，可能是扫描件或图片型PDF，请手动输入合同信息】"
        return text
    elif ext in ['.doc', '.docx']:
        try:
            # 先尝试使用 python-docx (推荐方式)
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # 如果段落为空，尝试从表格中提取
            if not text_parts:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text.strip())
            content = '\n'.join(text_parts)
            if not content.strip():
                return "【DOCX文件中未找到可提取的文本内容】"
            return content
        except ImportError:
            # 如果 python-docx 未安装，尝试使用 antiword 或其他方式
            try:
                import subprocess
                result = subprocess.run(['antiword', file_path], 
                                       capture_output=True, text=True, timeout=30)
                if result.returncode == 0:
                    return result.stdout
            except:
                pass
            return "【DOCX文件需要安装 python-docx 库才能解析，请运行: pip install python-docx】"
        except Exception as e:
            # 记录详细错误以便调试
            print(f"[ERROR] Failed to read docx file {file_path}: {e}")
            return f"【无法解析DOCX文件: {str(e)}，请手动输入合同信息】"
    elif ext in ['.jpg', '.jpeg', '.png']:
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(file_path)
            return pytesseract.image_to_string(img, lang='chi_sim')
        except ImportError:
            return "【图片文件需要安装pytesseract和Pillow库才能解析，请手动输入合同信息】"
        except Exception:
            return "【无法解析图片内容，请手动输入合同信息】"
    else:
        return "【不支持的文件格式，请手动输入合同信息】"


# ==================== 路由：前端页面 ====================

@app.route('/')
def home():
    """首页 - 基金列表"""
    return render_template('index.html')


@app.route('/detail/<int:fund_id>')
def fund_detail(fund_id):
    """基金详情页"""
    return render_template('detail.html')


@app.route('/static/<path:filename>')
def serve_static(filename):
    """提供静态文件"""
    return send_from_directory('static', filename)


# ==================== 路由：根路径 ====================

@app.route('/api')
def api_info():
    """API 根路径，返回欢迎信息"""
    return jsonify({
        'message': 'Welcome to Fund Management System API',
        'version': '1.0.0',
        'endpoints': {
            'funds': '/api/funds',
            'nodes': '/api/nodes/<id>',
            'upload': '/api/nodes/<id>/upload'
        }
    })


# ==================== 路由：基金 (Funds) ====================

@app.route('/api/stats', methods=['GET'])
def get_stats():
    """
    首页顶部统计：基金总数、节点总数、文件总数、今日新增
    """
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute('SELECT COUNT(*) AS c FROM funds')
    total_funds = cursor.fetchone()['c']

    cursor.execute('SELECT COUNT(*) AS c FROM nodes')
    total_nodes = cursor.fetchone()['c']

    cursor.execute('SELECT COUNT(*) AS c FROM files')
    total_files = cursor.fetchone()['c']

    # 今日新增（按北京时间今日 0 点 - 24 点）
    cursor.execute('''
        SELECT COUNT(*) AS c FROM funds
        WHERE created_at >= datetime('now', 'start of day', '+0 hours')
          AND created_at <  datetime('now', 'start of day', '+1 day', '+0 hours')
    ''')
    today_funds = cursor.fetchone()['c']

    conn.close()
    return jsonify({
        'code': 200,
        'message': 'success',
        'data': {
            'total_funds': total_funds,
            'today_funds': today_funds,
            'total_nodes': total_nodes,
            'total_files': total_files
        }
    })


@app.route('/api/funds', methods=['GET'])
def get_funds():
    """
    获取所有基金列表
    支持按名称或代码搜索 (?search=关键词)
    """
    conn = get_connection()
    cursor = conn.cursor()
    
    search = request.args.get('search', '')
    
    if search:
        cursor.execute('''
            SELECT * FROM funds 
            WHERE name LIKE ? OR code LIKE ?
            ORDER BY created_at DESC
        ''', (f'%{search}%', f'%{search}%'))
    else:
        cursor.execute('SELECT * FROM funds ORDER BY created_at DESC')
    
    funds = [dict_from_row(row) for row in cursor.fetchall()]
    conn.close()
    
    return jsonify({
        'code': 200,
        'message': 'success',
        'data': funds
    })


@app.route('/api/funds/<int:fund_id>', methods=['GET'])
def get_fund(fund_id):
    """获取单个基金详情"""
    conn = get_connection()
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM funds WHERE id = ?', (fund_id,))
    fund = dict_from_row(cursor.fetchone())
    
    if not fund:
        conn.close()
        return jsonify({'code': 404, 'message': '基金不存在'}), 404
    
    conn.close()
    return jsonify({
        'code': 200,
        'message': 'success',
        'data': fund
    })


@app.route('/api/funds', methods=['POST'])
def create_fund():
    """创建新基金"""
    data = request.get_json()
    
    if not data:
        return jsonify({'code': 400, 'message': '请求数据不能为空'}), 400
    
    name = data.get('name', '').strip()
    code = data.get('code', '').strip()
    
    if not name or not code:
        return jsonify({'code': 400, 'message': '基金名称和代码不能为空'}), 400
    
    conn = get_connection()
    cursor = conn.cursor()
    
    # 检查代码是否已存在
    cursor.execute('SELECT id FROM funds WHERE code = ?', (code,))
    if cursor.fetchone():
        conn.close()
        return jsonify({'code': 400, 'message': '基金代码已存在'}), 400
    
    try:
        cursor.execute('''
            INSERT INTO funds (name, code, created_at, updated_at)
            VALUES (?, ?, ?, ?)
        ''', (name, code, datetime.now(), datetime.now()))
        
        fund_id = cursor.lastrowid
        conn.commit()
        
        cursor.execute('SELECT * FROM funds WHERE id = ?', (fund_id,))
        fund = dict_from_row(cursor.fetchone())
        conn.close()
        
        return jsonify({
            'code': 200,
            'message': '基金创建成功',
            'data': fund
        })
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({'code': 500, 'message': f'创建失败: {str(e)}'}), 500


@app.route('/api/funds/<int:fund_id>', methods=['PUT'])
def update_fund(fund_id):
    """更新基金信息"""
    data = request.get_json()
    
    if not data:
        return jsonify({'code': 400, 'message': '请求数据不能为空'}), 400
    
    conn = get_connection()
    cursor = conn.cursor()
    
    # 检查基金是否存在
    cursor.execute('SELECT * FROM funds WHERE id = ?', (fund_id,))
    if not cursor.fetchone():
        conn.close()
        return jsonify({'code': 404, 'message': '基金不存在'}), 404
    
    name = data.get('name', '').strip()
    code = data.get('code', '').strip()
    
    if not name or not code:
        conn.close()
        return jsonify({'code': 400, 'message': '基金名称和代码不能为空'}), 400
    
    try:
        cursor.execute('''
            UPDATE funds 
            SET name = ?, code = ?, updated_at = ?
            WHERE id = ?
        ''', (name, code, datetime.now(), fund_id))
        
        conn.commit()
        
        cursor.execute('SELECT * FROM funds WHERE id = ?', (fund_id,))
        fund = dict_from_row(cursor.fetchone())
        conn.close()
        
        return jsonify({
            'code': 200,
            'message': '基金更新成功',
            'data': fund
        })
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({'code': 500, 'message': f'更新失败: {str(e)}'}), 500


@app.route('/api/funds/<int:fund_id>', methods=['DELETE'])
def delete_fund(fund_id):
    """删除基金（级联删除节点）"""
    conn = get_connection()
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM funds WHERE id = ?', (fund_id,))
    if not cursor.fetchone():
        conn.close()
        return jsonify({'code': 404, 'message': '基金不存在'}), 404
    
    try:
        cursor.execute('DELETE FROM funds WHERE id = ?', (fund_id,))
        conn.commit()
        conn.close()
        
        return jsonify({
            'code': 200,
            'message': '基金删除成功'
        })
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({'code': 500, 'message': f'删除失败: {str(e)}'}), 500


# ==================== 路由：节点 (Nodes) ====================

@app.route('/api/funds/<int:fund_id>/nodes', methods=['GET'])
def get_nodes_by_fund(fund_id):
    """获取基金的所有节点"""
    conn = get_connection()
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT n.*, ne.id as extraction_id
        FROM nodes n
        LEFT JOIN node_extractions ne ON n.id = ne.node_id
        WHERE n.fund_id = ?
        ORDER BY n.created_at DESC
    ''', (fund_id,))
    
    nodes = [dict_from_row(row) for row in cursor.fetchall()]
    conn.close()
    
    return jsonify({
        'code': 200,
        'message': 'success',
        'data': nodes
    })


@app.route('/api/nodes/<int:node_id>', methods=['GET'])
def get_node(node_id):
    """获取节点详情"""
    conn = get_connection()
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT n.*, f.name as fund_name, f.code as fund_code
        FROM nodes n
        JOIN funds f ON n.fund_id = f.id
        WHERE n.id = ?
    ''', (node_id,))
    node = dict_from_row(cursor.fetchone())
    
    if not node:
        conn.close()
        return jsonify({'code': 404, 'message': '节点不存在'}), 404
    
    # 获取提取信息
    cursor.execute('SELECT * FROM node_extractions WHERE node_id = ?', (node_id,))
    extraction = dict_from_row(cursor.fetchone())
    
    # 获取文件列表
    cursor.execute('SELECT * FROM files WHERE node_id = ?', (node_id,))
    files = [dict_from_row(row) for row in cursor.fetchall()]
    
    node['extraction'] = extraction
    node['files'] = files
    
    conn.close()
    
    return jsonify({
        'code': 200,
        'message': 'success',
        'data': node
    })


@app.route('/api/funds/<int:fund_id>/nodes', methods=['POST'])
def create_node(fund_id):
    """为基金创建新节点"""
    data = request.get_json()
    
    if not data:
        return jsonify({'code': 400, 'message': '请求数据不能为空'}), 400
    
    node_type = data.get('node_type', '').strip()
    
    if not node_type:
        return jsonify({'code': 400, 'message': '节点类型不能为空'}), 400
    
    # 验证节点类型
    valid_types = ['首次', '变更', '分红', '临开', '自定义']
    if node_type not in valid_types:
        return jsonify({'code': 400, 'message': f'节点类型必须是: {", ".join(valid_types)}'}), 400
    
    conn = get_connection()
    cursor = conn.cursor()
    
    # 检查基金是否存在
    cursor.execute('SELECT id FROM funds WHERE id = ?', (fund_id,))
    if not cursor.fetchone():
        conn.close()
        return jsonify({'code': 404, 'message': '基金不存在'}), 404
    
    try:
        cursor.execute('''
            INSERT INTO nodes (fund_id, node_type, created_at, updated_at)
            VALUES (?, ?, ?, ?)
        ''', (fund_id, node_type, datetime.now(), datetime.now()))
        
        node_id = cursor.lastrowid
        conn.commit()
        
        cursor.execute('SELECT * FROM nodes WHERE id = ?', (node_id,))
        node = dict_from_row(cursor.fetchone())
        conn.close()
        
        return jsonify({
            'code': 200,
            'message': '节点创建成功',
            'data': node
        })
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({'code': 500, 'message': f'创建失败: {str(e)}'}), 500


@app.route('/api/nodes/<int:node_id>', methods=['DELETE'])
def delete_node(node_id):
    """删除节点"""
    conn = get_connection()
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM nodes WHERE id = ?', (node_id,))
    if not cursor.fetchone():
        conn.close()
        return jsonify({'code': 404, 'message': '节点不存在'}), 404
    
    try:
        cursor.execute('DELETE FROM nodes WHERE id = ?', (node_id,))
        conn.commit()
        conn.close()
        
        return jsonify({
            'code': 200,
            'message': '节点删除成功'
        })
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({'code': 500, 'message': f'删除失败: {str(e)}'}), 500


# ==================== 路由：提取信息 (Extractions) ====================

@app.route('/api/nodes/<int:node_id>/ai-extract', methods=['POST'])
def ai_extract(node_id):
    """
    使用 DeepSeek AI 提取节点文件中的信息
    """
    conn = get_connection()
    cursor = conn.cursor()
    
    # 检查节点是否存在
    cursor.execute('SELECT * FROM nodes WHERE id = ?', (node_id,))
    node = cursor.fetchone()
    if not node:
        conn.close()
        return jsonify({'code': 404, 'message': '节点不存在'}), 404
    
    node_type = dict_from_row(node)['node_type']
    
    # 获取节点的所有文件
    cursor.execute('SELECT * FROM files WHERE node_id = ? ORDER BY upload_time DESC', (node_id,))
    files = [dict_from_row(row) for row in cursor.fetchall()]
    conn.close()
    
    if not files:
        return jsonify({'code': 400, 'message': '请先上传文件再进行提取'}), 400
    
    try:
        # 读取第一个文件的内容
        file_record = files[0]
        print(f"[INFO] Reading file: {file_record['file_path']}")
        content = read_file_content(file_record['file_path'])
        print(f"[INFO] File content length: {len(content) if content else 0}")
        
        if not content or content.startswith('【'):
            print(f"[ERROR] Failed to read file content: {content}")
            return jsonify({
                'code': 400,
                'message': content or '无法读取文件内容',
                'data': None
            }), 400
        
        # 调用 DeepSeek 提取
        print(f"[INFO] Calling DeepSeek API for node type: {node_type}")
        extracted = extract_with_deepseek(content, node_type)
        print(f"[INFO] DeepSeek returned: {extracted}")
        
        if not extracted:
            return jsonify({
                'code': 500,
                'message': 'AI提取失败，请稍后重试',
                'data': None
            }), 500
        
        # 保存提取结果到数据库
        save_extraction(node_id, extracted)
        
        return jsonify({
            'code': 200,
            'message': '提取成功',
            'data': extracted
        })
        
    except ValueError as e:
        # 业务逻辑错误（如内容为空、API错误等）
        print(f"[ERROR] Business error: {e}")
        return jsonify({
            'code': 400,
            'message': str(e),
            'data': None
        }), 400
    except Exception as e:
        # 其他未知错误
        print(f"[ERROR] Unexpected error in ai_extract: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'code': 500,
            'message': f'提取失败: {str(e)}',
            'data': None
        }), 500


def save_extraction(node_id, data):
    """保存提取信息到数据库"""
    conn = get_connection()
    cursor = conn.cursor()
    
    # 检查是否已有提取信息
    cursor.execute('SELECT id FROM node_extractions WHERE node_id = ?', (node_id,))
    existing = cursor.fetchone()
    
    if existing:
        # 更新现有记录
        fields = []
        values = []
        for key in ['product_name', 'custodian', 'establish_date', 'investment_manager',
                   'strategy', 'subscription_fee', 'custody_fee', 'risk_return',
                   'lock_period', 'annual_fee', 'performance_fee', 'investment_scope',
                   'investment_ratio', 'dividend_count', 'dividend_date', 
                   'dividend_amount', 'remark']:
            if key in data:
                fields.append(f'{key} = ?')
                values.append(data[key])
        
        fields.append('updated_at = ?')
        values.append(datetime.now())
        values.append(node_id)
        
        sql = f"UPDATE node_extractions SET {', '.join(fields)} WHERE node_id = ?"
        cursor.execute(sql, values)
    else:
        # 创建新记录
        cursor.execute('''
            INSERT INTO node_extractions (node_id, product_name, custodian, establish_date,
                investment_manager, strategy, subscription_fee, custody_fee, risk_return,
                lock_period, annual_fee, performance_fee, investment_scope, investment_ratio,
                dividend_count, dividend_date, dividend_amount, remark, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            node_id, data.get('product_name'), data.get('custodian'),
            data.get('establish_date'), data.get('investment_manager'),
            data.get('strategy'), data.get('subscription_fee'), data.get('custody_fee'),
            data.get('risk_return'), data.get('lock_period'), data.get('annual_fee'),
            data.get('performance_fee'), data.get('investment_scope'), data.get('investment_ratio'),
            data.get('dividend_count'), data.get('dividend_date'), data.get('dividend_amount'),
            data.get('remark'), datetime.now(), datetime.now()
        ))
    
    conn.commit()
    conn.close()


@app.route('/api/nodes/<int:node_id>/extraction', methods=['GET'])
def get_extraction(node_id):
    """获取节点提取信息"""
    conn = get_connection()
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM node_extractions WHERE node_id = ?', (node_id,))
    extraction = dict_from_row(cursor.fetchone())
    
    conn.close()
    
    if not extraction:
        return jsonify({
            'code': 200,
            'message': 'success',
            'data': None
        })
    
    return jsonify({
        'code': 200,
        'message': 'success',
        'data': extraction
    })


@app.route('/api/nodes/<int:node_id>/extraction', methods=['POST'])
def create_or_update_extraction(node_id):
    """创建或更新节点提取信息"""
    data = request.get_json()
    
    if not data:
        return jsonify({'code': 400, 'message': '请求数据不能为空'}), 400
    
    conn = get_connection()
    cursor = conn.cursor()
    
    # 检查节点是否存在
    cursor.execute('SELECT * FROM nodes WHERE id = ?', (node_id,))
    if not cursor.fetchone():
        conn.close()
        return jsonify({'code': 404, 'message': '节点不存在'}), 404
    
    # 检查是否已有提取信息
    cursor.execute('SELECT id FROM node_extractions WHERE node_id = ?', (node_id,))
    existing = cursor.fetchone()
    
    try:
        if existing:
            # 更新现有记录
            fields = []
            values = []
            
            for key in ['product_name', 'custodian', 'establish_date', 'investment_manager',
                       'strategy', 'subscription_fee', 'custody_fee', 'risk_return',
                       'lock_period', 'annual_fee', 'performance_fee', 'investment_scope',
                       'investment_ratio', 'dividend_count', 'dividend_date', 
                       'dividend_amount', 'remark']:
                if key in data:
                    fields.append(f'{key} = ?')
                    values.append(data[key])
            
            fields.append('updated_at = ?')
            values.append(datetime.now())
            values.append(node_id)
            
            sql = f"UPDATE node_extractions SET {', '.join(fields)} WHERE node_id = ?"
            cursor.execute(sql, values)
        else:
            # 创建新记录
            cursor.execute('''
                INSERT INTO node_extractions (
                    node_id, product_name, custodian, establish_date,
                    investment_manager, strategy, subscription_fee, custody_fee,
                    risk_return, lock_period, annual_fee, performance_fee,
                    investment_scope, investment_ratio, dividend_count,
                    dividend_date, dividend_amount, remark, created_at, updated_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                node_id,
                data.get('product_name'),
                data.get('custodian'),
                data.get('establish_date'),
                data.get('investment_manager'),
                data.get('strategy'),
                data.get('subscription_fee'),
                data.get('custody_fee'),
                data.get('risk_return'),
                data.get('lock_period'),
                data.get('annual_fee'),
                data.get('performance_fee'),
                data.get('investment_scope'),
                data.get('investment_ratio'),
                data.get('dividend_count'),
                data.get('dividend_date'),
                data.get('dividend_amount'),
                data.get('remark'),
                datetime.now(),
                datetime.now()
            ))
        
        conn.commit()
        
        cursor.execute('SELECT * FROM node_extractions WHERE node_id = ?', (node_id,))
        extraction = dict_from_row(cursor.fetchone())
        conn.close()
        
        return jsonify({
            'code': 200,
            'message': '提取信息保存成功',
            'data': extraction
        })
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({'code': 500, 'message': f'保存失败: {str(e)}'}), 500


# ==================== 路由：文件上传 (Files) ====================

def allowed_file(filename):
    """检查文件扩展名是否允许"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


@app.route('/api/nodes/<int:node_id>/upload', methods=['POST'])
def upload_file(node_id):
    """上传文件到节点"""
    conn = get_connection()
    cursor = conn.cursor()
    
    # 检查节点是否存在
    cursor.execute('SELECT * FROM nodes WHERE id = ?', (node_id,))
    if not cursor.fetchone():
        conn.close()
        return jsonify({'code': 404, 'message': '节点不存在'}), 404
    
    # 检查是否有文件
    if 'file' not in request.files:
        conn.close()
        return jsonify({'code': 400, 'message': '没有文件'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        conn.close()
        return jsonify({'code': 400, 'message': '文件名为空'}), 400
    
    if not allowed_file(file.filename):
        conn.close()
        return jsonify({'code': 400, 'message': '不支持的文件类型'}), 400
    
    # 创建上传目录
    upload_dir = os.path.join(app.config['UPLOAD_FOLDER'], str(node_id))
    os.makedirs(upload_dir, exist_ok=True)
    
    # 保存文件
    filename = secure_filename(file.filename)
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    save_filename = f'{timestamp}_{filename}'
    file_path = os.path.join(upload_dir, save_filename)
    
    try:
        file.save(file_path)
        file_size = os.path.getsize(file_path)
        
        # 保存到数据库
        cursor.execute('''
            INSERT INTO files (node_id, filename, file_path, file_size, upload_time)
            VALUES (?, ?, ?, ?, ?)
        ''', (node_id, filename, file_path, file_size, datetime.now()))
        
        file_id = cursor.lastrowid
        conn.commit()
        
        cursor.execute('SELECT * FROM files WHERE id = ?', (file_id,))
        file_record = dict_from_row(cursor.fetchone())
        conn.close()
        
        return jsonify({
            'code': 200,
            'message': '文件上传成功',
            'data': file_record
        })
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({'code': 500, 'message': f'上传失败: {str(e)}'}), 500


@app.route('/api/nodes/<int:node_id>/files', methods=['GET'])
def get_files(node_id):
    """获取节点的所有文件"""
    conn = get_connection()
    cursor = conn.cursor()

    # 先检查节点是否存在
    cursor.execute('SELECT * FROM nodes WHERE id = ?', (node_id,))
    if not cursor.fetchone():
        conn.close()
        return jsonify({'code': 404, 'message': '节点不存在'}), 404

    cursor.execute('SELECT * FROM files WHERE node_id = ? ORDER BY upload_time DESC', (node_id,))
    files = [dict_from_row(row) for row in cursor.fetchall()]

    conn.close()

    return jsonify({
        'code': 200,
        'message': 'success',
        'data': files
    })


@app.route('/api/nodes/<int:node_id>/files/<int:file_id>', methods=['DELETE'])
def delete_file(node_id, file_id):
    """删除节点下的某个文件（从数据库与磁盘）"""
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute('SELECT * FROM files WHERE id = ? AND node_id = ?', (file_id, node_id))
    file_record = dict_from_row(cursor.fetchone())
    if not file_record:
        conn.close()
        return jsonify({'code': 404, 'message': '文件不存在'}), 404

    try:
        # 删除磁盘文件
        try:
            if file_record.get('file_path') and os.path.exists(file_record['file_path']):
                os.remove(file_record['file_path'])
        except Exception as e:
            print(f"[WARNING] 删除磁盘文件失败: {e}")

        cursor.execute('DELETE FROM files WHERE id = ?', (file_id,))
        conn.commit()
        conn.close()
        return jsonify({'code': 200, 'message': '文件删除成功'})
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({'code': 500, 'message': f'删除失败: {str(e)}'}), 500


@app.route('/api/nodes/<int:node_id>/files/<int:file_id>/download', methods=['GET'])
def download_file(node_id, file_id):
    """下载节点下的某个文件"""
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute('SELECT * FROM files WHERE id = ? AND node_id = ?', (file_id, node_id))
    file_record = dict_from_row(cursor.fetchone())
    if not file_record:
        conn.close()
        return jsonify({'code': 404, 'message': '文件不存在'}), 404
    conn.close()

    # send_from_directory 需要 directory + path
    directory = os.path.dirname(file_record['file_path'])
    filename = os.path.basename(file_record['file_path'])
    return send_from_directory(directory, filename, as_attachment=True, download_name=file_record['filename'])


# ==================== 错误处理 ====================

@app.errorhandler(404)
def not_found(error):
    return jsonify({'code': 404, 'message': '接口不存在'}), 404


@app.errorhandler(500)
def internal_error(error):
    return jsonify({'code': 500, 'message': '服务器内部错误'}), 500


# ==================== 启动应用 ====================

if __name__ == '__main__':
    # 检查数据库是否存在
    if not os.path.exists(DATABASE):
        print(f"[ERROR] Database '{DATABASE}' not found!")
        print("Please run 'python init_db.py' first.")
        exit(1)
    
    # 创建上传目录
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    
    print("="*60)
    print("FUND MANAGEMENT SYSTEM - Flask API Server")
    print("="*60)
    print(f"Database: {os.path.abspath(DATABASE)}")
    print(f"Upload folder: {os.path.abspath(app.config['UPLOAD_FOLDER'])}")
    print("="*60)
    print("Server running at: http://127.0.0.1:5000")
    print("API Documentation:")
    print("  GET    /api/funds              - Get all funds")
    print("  GET    /api/funds/<id>         - Get fund by ID")
    print("  POST   /api/funds              - Create fund")
    print("  PUT    /api/funds/<id>         - Update fund")
    print("  DELETE /api/funds/<id>        - Delete fund")
    print("  GET    /api/funds/<id>/nodes   - Get nodes by fund")
    print("  POST   /api/funds/<id>/nodes   - Create node")
    print("  GET    /api/nodes/<id>        - Get node details")
    print("  DELETE /api/nodes/<id>        - Delete node")
    print("  GET    /api/nodes/<id>/extraction  - Get extraction")
    print("  POST   /api/nodes/<id>/extraction  - Save extraction")
    print("  POST   /api/nodes/<id>/upload      - Upload file")
    print("  GET    /api/nodes/<id>/files       - Get files")
    print("="*60)
    
    # 启动服务器，debug模式
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)), debug=True)
